import os
import re
import json
import shutil
import requests
from datetime import datetime
from typing import Optional

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    CallbackQueryHandler,
    ContextTypes,
    filters,
)
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit

# =========================
# SOZLAMALAR
# =========================
TOKEN = "8290808043:AAEWX_oDghVc0teDOgJbYVNjnnLXqXdCYmg"
ADMIN_ID = 7013209195
SUPPORT_USERNAME = "@mirshod_choriyev"

# DeepSeek AI
DEEPSEEK_API_KEY = "sk-8e28a062b1204b84bba6b7d98e9a1dae"
DEEPSEEK_URL = "https://api.deepseek.com/v1/chat/completions"
DEEPSEEK_MODEL = "deepseek-chat"

DATA_FILE = "stats.json"

# =========================
# XOTIRA
# =========================
user_images = {}
user_pdf_sizes = {}
user_states = {}

# =========================
# MATNLAR
# =========================
TXT = {
    "welcome": (
        "Assalomu alaykum 👋\n\n"
        "Men sizga:\n"
        "• rasmlarni PDF/Word ga aylantiraman\n"
        "• matnni chiroyli Word/PDF qilaman\n"
        "• AI orqali matn yozib beraman\n\n"
        "Kerakli tugmani tanlang 👇"
    ),
    "photo_saved": "✅ {count} ta rasm saqlandi.",
    "no_images": "❌ Avval rasm yuboring.",
    "ask_filename_pdf": "📄 PDF fayl nomini yuboring.\nMasalan: fizika_ishi",
    "ask_filename_word": "📝 Word fayl nomini yuboring.\nMasalan: fizika_ishi",
    "ask_text_pdf": "📄 PDF ga tushirish uchun matn yuboring.",
    "ask_text_word": "📝 Word ga tushirish uchun matn yuboring.",
    "ask_ai_text": "🤖 AI yozib berishi uchun mavzu yoki matn yuboring.",
    "ask_ai_pdf": "🤖 AI chiroyli PDF qilish uchun matn/mavzu yuboring.",
    "ask_ai_word": "🤖 AI chiroyli Word qilish uchun matn/mavzu yuboring.",
    "pdf_size_saved": "✅ PDF hajmi saqlandi.",
    "cleared": "🧹 Fayllar tozalandi.",
    "loading_photo": "⏳ Rasm saqlanmoqda...",
    "loading_pdf": "⏳ PDF tayyorlanmoqda...",
    "loading_word": "⏳ Word tayyorlanmoqda...",
    "loading_ai": "🤖 AI ishlayapti...",
    "done_pdf": "✅ PDF tayyor.",
    "done_word": "✅ Word tayyor.",
    "done_ai": "✅ AI javobi tayyor.",
    "invalid_name": "❌ Fayl nomi noto‘g‘ri.",
    "long_name": "❌ Fayl nomi juda uzun.",
    "help": (
        "🆘 Yordam xizmati\n\n"
        "Muammo bo‘lsa menga yozing:\n"
        f"{SUPPORT_USERNAME}"
    ),
    "not_admin": "⛔ Siz admin emassiz.",
    "admin_title": "👑 Admin panel",
    "general_error": "❌ Xatolik yuz berdi. Qayta urinib ko‘ring.",
    "myid_text": "🆔 Sizning Telegram ID: {id}",
    "ai_not_configured": "❌ AI hali sozlanmagan.",
}

# =========================
# STATISTIKA
# =========================
def load_stats():
    if not os.path.exists(DATA_FILE):
        return {
            "users": {},
            "totals": {
                "image_count": 0,
                "pdf_count": 0,
                "word_count": 0,
                "text_pdf_count": 0,
                "text_word_count": 0,
                "ai_count": 0,
            },
        }
    with open(DATA_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def save_stats(data):
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


stats_data = load_stats()


def ensure_user_exists(user_id: int):
    uid = str(user_id)
    if uid not in stats_data["users"]:
        stats_data["users"][uid] = {
            "first_seen": datetime.now().strftime("%Y-%m-%d"),
            "last_seen": datetime.now().strftime("%Y-%m-%d"),
            "daily_visits": [],
            "monthly_visits": [],
        }


def track_user(user_id: int):
    today = datetime.now().strftime("%Y-%m-%d")
    month = datetime.now().strftime("%Y-%m")

    ensure_user_exists(user_id)
    uid = str(user_id)

    stats_data["users"][uid]["last_seen"] = today

    if today not in stats_data["users"][uid]["daily_visits"]:
        stats_data["users"][uid]["daily_visits"].append(today)

    if month not in stats_data["users"][uid]["monthly_visits"]:
        stats_data["users"][uid]["monthly_visits"].append(month)

    save_stats(stats_data)


def get_total_users():
    return len(stats_data["users"])


def get_today_users():
    today = datetime.now().strftime("%Y-%m-%d")
    count = 0
    for user in stats_data["users"].values():
        if today in user.get("daily_visits", []):
            count += 1
    return count


def get_month_users():
    month = datetime.now().strftime("%Y-%m")
    count = 0
    for user in stats_data["users"].values():
        if month in user.get("monthly_visits", []):
            count += 1
    return count


# =========================
# FAYL YORDAMCHILARI
# =========================
def get_user_folder(user_id: int) -> str:
    folder = f"files/{user_id}"
    os.makedirs(folder, exist_ok=True)
    return folder


def get_temp_folder(user_id: int) -> str:
    folder = f"files/{user_id}/temp"
    os.makedirs(folder, exist_ok=True)
    return folder


def cleanup_temp_folder(user_id: int):
    temp_folder = f"files/{user_id}/temp"
    if os.path.exists(temp_folder):
        shutil.rmtree(temp_folder, ignore_errors=True)


def sanitize_filename(name: str) -> str:
    name = name.strip()
    name = re.sub(r'[\\/:*?"<>|]+', "", name)
    name = re.sub(r"\s+", "_", name)
    return name


def unique_path(base_folder: str, file_name: str, ext: str) -> str:
    path = os.path.join(base_folder, f"{file_name}.{ext}")
    if not os.path.exists(path):
        return path

    i = 2
    while True:
        path = os.path.join(base_folder, f"{file_name}_{i}.{ext}")
        if not os.path.exists(path):
            return path
        i += 1


def get_pdf_quality(size_choice: str) -> int:
    if size_choice == "1":
        return 20
    if size_choice == "2":
        return 35
    if size_choice == "5":
        return 55
    return 95


def prepare_images_for_pdf(user_id: int, size_choice: str):
    original_paths = user_images.get(user_id, [])
    temp_paths = []

    if size_choice == "org":
        return original_paths

    quality = get_pdf_quality(size_choice)
    temp_folder = get_temp_folder(user_id)

    for i, path in enumerate(original_paths, start=1):
        img = Image.open(path).convert("RGB")
        temp_path = os.path.join(temp_folder, f"compressed_{i}.jpg")
        img.save(temp_path, "JPEG", quality=quality, optimize=True)
        img.close()
        temp_paths.append(temp_path)

    return temp_paths


# =========================
# AI
# =========================
def ai_ready() -> bool:
    return (
        DEEPSEEK_API_KEY
        and DEEPSEEK_API_KEY != "YOUR_DEEPSEEK_API_KEY"
        and DEEPSEEK_MODEL
    )


def ai_generate_text(user_text: str) -> Optional[str]:
    if not ai_ready():
        return None

    response = requests.post(
        DEEPSEEK_URL,
        headers={
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json",
        },
        json={
            "model": DEEPSEEK_MODEL,
            "messages": [
                {
                    "role": "system",
                    "content": "Siz foydali, chiroyli va aniq yozadigan yordamchisiz.",
                },
                {
                    "role": "user",
                    "content": (
                        "Quyidagi matn yoki mavzuni chiroyli, tartibli, savodli va foydali qilib yozib ber. "
                        "Kerak bo‘lsa sarlavha qo‘y, punktlar bilan bezat, natija hujjatga tayyor bo‘lsin.\n\n"
                        f"{user_text}"
                    ),
                },
            ],
            "temperature": 0.7,
        },
        timeout=90,
    )
    response.raise_for_status()
    data = response.json()
    return data["choices"][0]["message"]["content"].strip()


# =========================
# MATN -> WORD/PDF
# =========================
def create_word_from_text(file_path: str, title: str, content: str):
    doc = Document()

    title_p = doc.add_paragraph()
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph("")

    for part in content.split("\n"):
        doc.add_paragraph(part)

    doc.save(file_path)


def create_pdf_from_text(file_path: str, title: str, content: str):
    c = canvas.Canvas(file_path, pagesize=A4)
    width, height = A4
    x_margin = 50
    y = height - 60

    c.setFont("Helvetica-Bold", 16)
    c.drawString(x_margin, y, title)
    y -= 30

    c.setFont("Helvetica", 11)
    lines = []

    for paragraph in content.split("\n"):
        wrapped = simpleSplit(paragraph, "Helvetica", 11, width - 2 * x_margin)
        if not wrapped:
            lines.append("")
        else:
            lines.extend(wrapped)

    for line in lines:
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - 50
        c.drawString(x_margin, y, line)
        y -= 16

    c.save()


# =========================
# MENYULAR
# =========================
def get_main_menu(user_id: int):
    keyboard = [
        [
            InlineKeyboardButton("📄 PDF qilish", callback_data="img_to_pdf"),
            InlineKeyboardButton("📝 Word qilish", callback_data="img_to_word"),
        ],
        [
            InlineKeyboardButton("📉 PDF hajmi", callback_data="choose_size"),
            InlineKeyboardButton("🧹 Tozalash", callback_data="clear_files"),
        ],
        [
            InlineKeyboardButton("🤖 AI panel", callback_data="ai_panel"),
            InlineKeyboardButton("📝 Matndan fayl", callback_data="text_panel"),
        ],
        [
            InlineKeyboardButton("🆘 Yordam", callback_data="help_contact"),
        ],
    ]

    if user_id == ADMIN_ID:
        keyboard.append([InlineKeyboardButton("👑 Admin panel", callback_data="admin_panel")])

    return InlineKeyboardMarkup(keyboard)


def get_pdf_size_menu():
    keyboard = [
        [
            InlineKeyboardButton("1 MB", callback_data="size_1"),
            InlineKeyboardButton("2 MB", callback_data="size_2"),
        ],
        [
            InlineKeyboardButton("5 MB", callback_data="size_5"),
            InlineKeyboardButton("Original", callback_data="size_org"),
        ],
        [
            InlineKeyboardButton("⬅️ Orqaga", callback_data="back_main"),
        ],
    ]
    return InlineKeyboardMarkup(keyboard)


def get_ai_menu():
    keyboard = [
        [InlineKeyboardButton("✨ AI matn yozsin", callback_data="ai_text_only")],
        [
            InlineKeyboardButton("🤖 AI → PDF", callback_data="ai_to_pdf"),
            InlineKeyboardButton("🤖 AI → Word", callback_data="ai_to_word"),
        ],
        [
            InlineKeyboardButton("⬅️ Orqaga", callback_data="back_main"),
        ],
    ]
    return InlineKeyboardMarkup(keyboard)


def get_text_menu():
    keyboard = [
        [
            InlineKeyboardButton("📄 Matndan PDF", callback_data="text_to_pdf"),
            InlineKeyboardButton("📝 Matndan Word", callback_data="text_to_word"),
        ],
        [
            InlineKeyboardButton("⬅️ Orqaga", callback_data="back_main"),
        ],
    ]
    return InlineKeyboardMarkup(keyboard)


def get_admin_menu():
    keyboard = [
        [
            InlineKeyboardButton("📊 Statistika", callback_data="admin_stats"),
            InlineKeyboardButton("🌍 Jami user", callback_data="admin_total_users"),
        ],
        [
            InlineKeyboardButton("👥 Bugungi user", callback_data="admin_today_users"),
            InlineKeyboardButton("🗓 Oylik user", callback_data="admin_month_users"),
        ],
        [
            InlineKeyboardButton("🧹 Tozalash", callback_data="admin_clear_server"),
            InlineKeyboardButton("🆔 Mening ID", callback_data="admin_myid"),
        ],
        [
            InlineKeyboardButton("⬅️ Orqaga", callback_data="back_main"),
        ],
    ]
    return InlineKeyboardMarkup(keyboard)


# =========================
# KOMANDALAR
# =========================
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.message:
        return
    user_id = update.message.from_user.id
    track_user(user_id)
    await update.message.reply_text(TXT["welcome"], reply_markup=get_main_menu(user_id))


async def myid(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.message:
        return
    await update.message.reply_text(TXT["myid_text"].format(id=update.message.from_user.id))


async def admin_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.message:
        return
    user_id = update.message.from_user.id
    if user_id != ADMIN_ID:
        await update.message.reply_text(TXT["not_admin"])
        return
    await update.message.reply_text(TXT["admin_title"], reply_markup=get_admin_menu())


async def stats_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.message:
        return
    user_id = update.message.from_user.id
    if user_id != ADMIN_ID:
        await update.message.reply_text(TXT["not_admin"])
        return
    await send_stats_text(update.message)


async def send_stats_text(message):
    text = (
        f"📊 Statistika\n\n"
        f"👥 Jami foydalanuvchilar: {get_total_users()}\n"
        f"📅 Bugungi foydalanuvchilar: {get_today_users()}\n"
        f"🗓 Oylik foydalanuvchilar: {get_month_users()}\n"
        f"🖼 Rasm soni: {stats_data['totals']['image_count']}\n"
        f"📄 PDF soni: {stats_data['totals']['pdf_count']}\n"
        f"📝 Word soni: {stats_data['totals']['word_count']}\n"
        f"📄 Matndan PDF: {stats_data['totals']['text_pdf_count']}\n"
        f"📝 Matndan Word: {stats_data['totals']['text_word_count']}\n"
        f"🤖 AI ishlatilgan: {stats_data['totals']['ai_count']}"
    )
    await message.reply_text(text, reply_markup=get_admin_menu())


# =========================
# RASM
# =========================
async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        if not update.message or not update.message.photo:
            return

        user_id = update.message.from_user.id
        track_user(user_id)

        folder = get_user_folder(user_id)
        user_images.setdefault(user_id, [])

        loading = await update.message.reply_text(TXT["loading_photo"])

        photo = update.message.photo[-1]
        tg_file = await context.bot.get_file(photo.file_id)

        file_path = os.path.join(folder, f"{photo.file_id}.jpg")
        await tg_file.download_to_drive(file_path)

        user_images[user_id].append(file_path)
        stats_data["totals"]["image_count"] += 1
        save_stats(stats_data)

        await loading.edit_text(
            TXT["photo_saved"].format(count=len(user_images[user_id])),
            reply_markup=get_main_menu(user_id),
        )
    except Exception:
        if update.message:
            await update.message.reply_text(TXT["general_error"])


# =========================
# IMAGE -> PDF/WORD
# =========================
async def make_pdf_from_images(message, user_id: int, file_name: str):
    if user_id not in user_images or not user_images[user_id]:
        await message.reply_text(TXT["no_images"])
        return

    try:
        loading = await message.reply_text(f"{TXT['loading_pdf']}\n\n0%")

        size_choice = user_pdf_sizes.get(user_id, "org")
        cleanup_temp_folder(user_id)
        image_paths = prepare_images_for_pdf(user_id, size_choice)

        await loading.edit_text(f"{TXT['loading_pdf']}\n\n35%")

        images = []
        for path in image_paths:
            img = Image.open(path).convert("RGB")
            images.append(img)

        await loading.edit_text(f"{TXT['loading_pdf']}\n\n75%")

        folder = get_user_folder(user_id)
        pdf_path = unique_path(folder, file_name, "pdf")
        images[0].save(pdf_path, save_all=True, append_images=images[1:])

        for img in images:
            img.close()

        stats_data["totals"]["pdf_count"] += 1
        save_stats(stats_data)

        await loading.edit_text(f"{TXT['loading_pdf']}\n\n100% ✅")

        with open(pdf_path, "rb") as f:
            await message.reply_document(
                f,
                filename=os.path.basename(pdf_path),
                caption=TXT["done_pdf"],
            )

        await loading.delete()
        cleanup_temp_folder(user_id)

    except Exception:
        await message.reply_text(TXT["general_error"])


async def make_word_from_images(message, user_id: int, file_name: str):
    if user_id not in user_images or not user_images[user_id]:
        await message.reply_text(TXT["no_images"])
        return

    try:
        loading = await message.reply_text(f"{TXT['loading_word']}\n\n0%")

        folder = get_user_folder(user_id)
        docx_path = unique_path(folder, file_name, "docx")
        doc = Document()

        await loading.edit_text(f"{TXT['loading_word']}\n\n45%")

        for path in user_images[user_id]:
            doc.add_picture(path, width=Inches(5.5))
            doc.add_paragraph("")

        await loading.edit_text(f"{TXT['loading_word']}\n\n85%")
        doc.save(docx_path)

        stats_data["totals"]["word_count"] += 1
        save_stats(stats_data)

        await loading.edit_text(f"{TXT['loading_word']}\n\n100% ✅")

        with open(docx_path, "rb") as f:
            await message.reply_document(
                f,
                filename=os.path.basename(docx_path),
                caption=TXT["done_word"],
            )

        await loading.delete()

    except Exception:
        await message.reply_text(TXT["general_error"])


# =========================
# TEXT -> PDF/WORD
# =========================
async def export_text_to_pdf(message, user_id: int, text_content: str, title="Matn hujjati"):
    try:
        loading = await message.reply_text(f"{TXT['loading_pdf']}\n\n0%")
        folder = get_user_folder(user_id)
        file_name = sanitize_filename(title) or "matn_hujjati"
        pdf_path = unique_path(folder, file_name, "pdf")

        await loading.edit_text(f"{TXT['loading_pdf']}\n\n50%")
        create_pdf_from_text(pdf_path, title, text_content)
        await loading.edit_text(f"{TXT['loading_pdf']}\n\n100% ✅")

        stats_data["totals"]["text_pdf_count"] += 1
        save_stats(stats_data)

        with open(pdf_path, "rb") as f:
            await message.reply_document(
                f,
                filename=os.path.basename(pdf_path),
                caption=TXT["done_pdf"],
            )

        await loading.delete()
    except Exception:
        await message.reply_text(TXT["general_error"])


async def export_text_to_word(message, user_id: int, text_content: str, title="Matn hujjati"):
    try:
        loading = await message.reply_text(f"{TXT['loading_word']}\n\n0%")
        folder = get_user_folder(user_id)
        file_name = sanitize_filename(title) or "matn_hujjati"
        docx_path = unique_path(folder, file_name, "docx")

        await loading.edit_text(f"{TXT['loading_word']}\n\n50%")
        create_word_from_text(docx_path, title, text_content)
        await loading.edit_text(f"{TXT['loading_word']}\n\n100% ✅")

        stats_data["totals"]["text_word_count"] += 1
        save_stats(stats_data)

        with open(docx_path, "rb") as f:
            await message.reply_document(
                f,
                filename=os.path.basename(docx_path),
                caption=TXT["done_word"],
            )

        await loading.delete()
    except Exception:
        await message.reply_text(TXT["general_error"])


# =========================
# MATN HANDLER
# =========================
async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.message or not update.message.text:
        return

    user_id = update.message.from_user.id
    track_user(user_id)
    mode = user_states.get(user_id)
    text_value = update.message.text.strip()

    if not mode:
        await update.message.reply_text("Kerakli tugmani tanlang 👇", reply_markup=get_main_menu(user_id))
        return

    if mode == "await_img_pdf_name":
        clean_name = sanitize_filename(text_value)
        if not clean_name:
            await update.message.reply_text(TXT["invalid_name"])
            return
        if len(clean_name) > 50:
            await update.message.reply_text(TXT["long_name"])
            return

        user_states.pop(user_id, None)
        await make_pdf_from_images(update.message, user_id, clean_name)
        return

    if mode == "await_img_word_name":
        clean_name = sanitize_filename(text_value)
        if not clean_name:
            await update.message.reply_text(TXT["invalid_name"])
            return
        if len(clean_name) > 50:
            await update.message.reply_text(TXT["long_name"])
            return

        user_states.pop(user_id, None)
        await make_word_from_images(update.message, user_id, clean_name)
        return

    if mode == "await_text_pdf":
        user_states.pop(user_id, None)
        title = text_value[:30] if len(text_value) > 5 else "matn_pdf"
        await export_text_to_pdf(update.message, user_id, text_value, title=title)
        return

    if mode == "await_text_word":
        user_states.pop(user_id, None)
        title = text_value[:30] if len(text_value) > 5 else "matn_word"
        await export_text_to_word(update.message, user_id, text_value, title=title)
        return

    if mode == "await_ai_text":
        user_states.pop(user_id, None)

        if not ai_ready():
            await update.message.reply_text(TXT["ai_not_configured"], reply_markup=get_main_menu(user_id))
            return

        try:
            loading = await update.message.reply_text(TXT["loading_ai"])
            result = ai_generate_text(text_value)
            stats_data["totals"]["ai_count"] += 1
            save_stats(stats_data)

            if not result:
                await loading.edit_text(TXT["ai_not_configured"])
                return

            await loading.edit_text("✅ AI javobi tayyor:\n\n" + result[:4000], reply_markup=get_ai_menu())
        except Exception:
            await update.message.reply_text(TXT["general_error"])
        return

    if mode == "await_ai_pdf":
        user_states.pop(user_id, None)

        if not ai_ready():
            await update.message.reply_text(TXT["ai_not_configured"], reply_markup=get_main_menu(user_id))
            return

        try:
            loading = await update.message.reply_text(TXT["loading_ai"])
            result = ai_generate_text(text_value)
            stats_data["totals"]["ai_count"] += 1
            save_stats(stats_data)

            if not result:
                await loading.edit_text(TXT["ai_not_configured"])
                return

            await loading.edit_text("🤖 AI matn tayyor. PDF qilinmoqda...\n\n50%")
            await export_text_to_pdf(update.message, user_id, result, title="ai_pdf_hujjat")
        except Exception:
            await update.message.reply_text(TXT["general_error"])
        return

    if mode == "await_ai_word":
        user_states.pop(user_id, None)

        if not ai_ready():
            await update.message.reply_text(TXT["ai_not_configured"], reply_markup=get_main_menu(user_id))
            return

        try:
            loading = await update.message.reply_text(TXT["loading_ai"])
            result = ai_generate_text(text_value)
            stats_data["totals"]["ai_count"] += 1
            save_stats(stats_data)

            if not result:
                await loading.edit_text(TXT["ai_not_configured"])
                return

            await loading.edit_text("🤖 AI matn tayyor. Word qilinmoqda...\n\n50%")
            await export_text_to_word(update.message, user_id, result, title="ai_word_hujjat")
        except Exception:
            await update.message.reply_text(TXT["general_error"])
        return

    await update.message.reply_text("Kerakli tugmani tanlang 👇", reply_markup=get_main_menu(user_id))


# =========================
# TUGMALAR
# =========================
async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user_id = query.from_user.id
    track_user(user_id)
    action = query.data

    try:
        if action == "img_to_pdf":
            if user_id not in user_images or not user_images[user_id]:
                await query.message.reply_text(TXT["no_images"])
                return
            user_states[user_id] = "await_img_pdf_name"
            await query.message.reply_text(TXT["ask_filename_pdf"])
            return

        if action == "img_to_word":
            if user_id not in user_images or not user_images[user_id]:
                await query.message.reply_text(TXT["no_images"])
                return
            user_states[user_id] = "await_img_word_name"
            await query.message.reply_text(TXT["ask_filename_word"])
            return

        if action == "choose_size":
            await query.message.reply_text("PDF hajmini tanlang 👇", reply_markup=get_pdf_size_menu())
            return

        if action.startswith("size_"):
            user_pdf_sizes[user_id] = action.split("_")[1]
            await query.message.reply_text(TXT["pdf_size_saved"], reply_markup=get_main_menu(user_id))
            return

        if action == "clear_files":
            user_images[user_id] = []
            user_states.pop(user_id, None)
            cleanup_temp_folder(user_id)

            folder = get_user_folder(user_id)
            if os.path.exists(folder):
                for file_name in os.listdir(folder):
                    file_path = os.path.join(folder, file_name)
                    if os.path.isfile(file_path):
                        os.remove(file_path)

            await query.message.reply_text(TXT["cleared"], reply_markup=get_main_menu(user_id))
            return

        if action == "ai_panel":
            await query.message.reply_text("🤖 AI panel", reply_markup=get_ai_menu())
            return

        if action == "text_panel":
            await query.message.reply_text("📝 Matndan fayl paneli", reply_markup=get_text_menu())
            return

        if action == "ai_text_only":
            user_states[user_id] = "await_ai_text"
            await query.message.reply_text(TXT["ask_ai_text"])
            return

        if action == "ai_to_pdf":
            user_states[user_id] = "await_ai_pdf"
            await query.message.reply_text(TXT["ask_ai_pdf"])
            return

        if action == "ai_to_word":
            user_states[user_id] = "await_ai_word"
            await query.message.reply_text(TXT["ask_ai_word"])
            return

        if action == "text_to_pdf":
            user_states[user_id] = "await_text_pdf"
            await query.message.reply_text(TXT["ask_text_pdf"])
            return

        if action == "text_to_word":
            user_states[user_id] = "await_text_word"
            await query.message.reply_text(TXT["ask_text_word"])
            return

        if action == "help_contact":
            await query.message.reply_text(TXT["help"], reply_markup=get_main_menu(user_id))
            return

        if action == "admin_panel":
            if user_id != ADMIN_ID:
                await query.message.reply_text(TXT["not_admin"])
                return
            await query.message.reply_text(TXT["admin_title"], reply_markup=get_admin_menu())
            return

        if action == "admin_stats":
            if user_id != ADMIN_ID:
                await query.message.reply_text(TXT["not_admin"])
                return
            await send_stats_text(query.message)
            return

        if action == "admin_total_users":
            if user_id != ADMIN_ID:
                await query.message.reply_text(TXT["not_admin"])
                return
            await query.message.reply_text(
                f"🌍 Jami foydalanuvchilar: {get_total_users()}",
                reply_markup=get_admin_menu(),
            )
            return

        if action == "admin_today_users":
            if user_id != ADMIN_ID:
                await query.message.reply_text(TXT["not_admin"])
                return
            await query.message.reply_text(
                f"👥 Bugungi foydalanuvchilar: {get_today_users()}",
                reply_markup=get_admin_menu(),
            )
            return

        if action == "admin_month_users":
            if user_id != ADMIN_ID:
                await query.message.reply_text(TXT["not_admin"])
                return
            await query.message.reply_text(
                f"🗓 Oylik foydalanuvchilar: {get_month_users()}",
                reply_markup=get_admin_menu(),
            )
            return

        if action == "admin_clear_server":
            if user_id != ADMIN_ID:
                await query.message.reply_text(TXT["not_admin"])
                return

            if os.path.exists("files"):
                shutil.rmtree("files", ignore_errors=True)
            os.makedirs("files", exist_ok=True)

            user_images.clear()
            await query.message.reply_text(
                "🧹 Serverdagi vaqtinchalik fayllar tozalandi.",
                reply_markup=get_admin_menu(),
            )
            return

        if action == "admin_myid":
            if user_id != ADMIN_ID:
                await query.message.reply_text(TXT["not_admin"])
                return
            await query.message.reply_text(
                TXT["myid_text"].format(id=user_id),
                reply_markup=get_admin_menu(),
            )
            return

        if action == "back_main":
            await query.message.reply_text("Asosiy menyu 👇", reply_markup=get_main_menu(user_id))
            return

    except Exception:
        await query.message.reply_text(TXT["general_error"], reply_markup=get_main_menu(user_id))


# =========================
# MAIN
# =========================
def main():
    if not TOKEN or TOKEN == "BOT_TOKEN_HERE":
        raise ValueError("TOKEN ni to‘g‘ri yozing")

    app = ApplicationBuilder().token(TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("myid", myid))
    app.add_handler(CommandHandler("admin", admin_command))
    app.add_handler(CommandHandler("stats", stats_command))

    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    app.add_handler(CallbackQueryHandler(button_handler))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))

    print("Bot ishga tushdi...")
    app.run_polling()


if __name__ == "__main__":
    main()